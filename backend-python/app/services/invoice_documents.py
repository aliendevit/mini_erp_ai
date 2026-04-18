from __future__ import annotations

from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path




SELLER = {
    "name": "Z & M Deco Bremen",
    "street": "Georg-Gleistein-Str. 7",
    "zipCity": "28757 Bremen",
    "phone": "016/21775100",
    "email": "Zyadkarkour@gmail.com",
}

FOOTER_LINES = [
    "Geschaeftsfuehrer: Zyad Karkour * Georg-Gleistein-Strasse 7 * 28757 Bremen * Tel: 01621775100",
    "Bankverbindung: Die Sparkasse Bremen BIC: SBREDE22XXX IBAN: DE06 2905 0101 0083 6017 08",
    "",
    "Freistellungsbescheinigung zu Steuerabzug bei Bauleistungen Paragraph 48 b Absatz 1 Satz 1 EStG liegt vor",
    "Finanzamt Bremen - Mitte, Steuernummer: 60/288/24293 - Ust-ID-Nr.: DE3592929773",
]


def _logo_path() -> Path | None:
    candidate = Path(__file__).resolve().parents[3] / "backend" / "assets" / "zmd-deco-logo.png"
    return candidate if candidate.exists() else None


def _safe_text(value: str | None) -> str:
    return (value or "").strip()


def _pick_issue_date(invoice) -> datetime:
    return invoice.issue_date or datetime.utcnow()


def _format_date_ddmmyy(value: datetime) -> str:
    return value.strftime("%d.%m.%y")


def _format_date_yyyymmdd(value: datetime) -> str:
    return value.strftime("%Y-%m-%d")


def _format_money_eur(amount: float) -> str:
    base = f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{base} EUR"


def _format_money_ui(amount: float) -> str:
    return f"{amount:.2f} EUR"


def _compute_line_rate(line) -> float:
    work_entry = line.work_entry
    if line.unit_rate is not None:
        return float(line.unit_rate)
    if work_entry and work_entry.order and work_entry.order.default_hourly_rate is not None:
        return float(work_entry.order.default_hourly_rate)
    if work_entry and work_entry.employee and work_entry.employee.default_hourly_rate is not None:
        return float(work_entry.employee.default_hourly_rate)
    return 0.0


def _compute_line_amount(line) -> float:
    if line.line_amount is not None:
        return float(line.line_amount)
    return float(line.hours_allocated or 0) * _compute_line_rate(line)


def _sorted_lines(invoice) -> list:
    return sorted(
        list(invoice.lines or []),
        key=lambda line: (
            line.service_date or datetime.min,
            str(line.id or ""),
        ),
    )


def _project_box_rows(invoice) -> tuple[str, str, str, str]:
    lines = list(invoice.lines or [])
    sites = []
    for line in lines:
        site_name = _safe_text(getattr(getattr(line.work_entry, "site", None), "site_name", None))
        if site_name and site_name not in sites:
            sites.append(site_name)
    site_label = ", ".join(sites) if sites else "-"

    service_dates = sorted([line.service_date for line in lines if line.service_date])
    if service_dates:
        period = f"{_format_date_ddmmyy(service_dates[0])}-{_format_date_ddmmyy(service_dates[-1])}"
    else:
        period = "-"

    order_desc = "-"
    for line in lines:
        description = _safe_text(getattr(getattr(line.work_entry, "order", None), "description", None))
        if description:
            order_desc = description
            break
    if order_desc == "-" and lines:
        order_desc = _safe_text(getattr(getattr(lines[0].work_entry, "order", None), "title", None)) or "-"

    customer = invoice.customer
    contact = _safe_text(getattr(customer, "contact_name", None)) or _safe_text(getattr(customer, "company_name", None)) or "-"
    bullet = f"* {order_desc}"
    fixed_price = f"Zum Festpreis laut Absprache vor Ort mit Herrn {contact}"
    return site_label, period, bullet, fixed_price


def _add_footer(canvas, _doc) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    width, height = A4
    y = 25 * mm
    canvas.saveState()
    canvas.setLineWidth(1)
    canvas.line(15 * mm, y, width - 15 * mm, y)
    canvas.setFont("Helvetica", 8)
    yy = y - 10
    for line in FOOTER_LINES:
        canvas.drawString(15 * mm, yy, line)
        yy -= 11
    canvas.restoreState()


def build_invoice_pdf(invoice, kind: str = "detailed") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=15 * mm, rightMargin=15 * mm, topMargin=15 * mm, bottomMargin=45 * mm)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("InvoiceNormal", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=12)
    bold = ParagraphStyle("InvoiceBold", parent=normal, fontName="Helvetica-Bold")
    small = ParagraphStyle("InvoiceSmall", parent=styles["Normal"], fontName="Helvetica", fontSize=8, leading=10)

    story = []

    logo = _logo_path()
    header_left = [Paragraph("<b>KUNDENRECHNUNG</b>", ParagraphStyle("Title", parent=bold, fontSize=16, leading=18))]
    header_right = []
    if logo:
        header_right.append(Image(str(logo), width=32 * mm, height=32 * mm))
    for key in ("name", "street", "zipCity", "phone", "email"):
        header_right.append(Paragraph(SELLER[key], normal))
    header = Table([[header_left, header_right]], colWidths=[100 * mm, 75 * mm])
    header.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header)
    story.append(Spacer(1, 6))

    customer = invoice.customer
    story.append(Paragraph(_safe_text(customer.company_name) or "-", normal))
    if _safe_text(customer.street):
        story.append(Paragraph(customer.street, normal))
    zip_city = " ".join([part for part in [_safe_text(customer.zip_code), _safe_text(customer.city)] if part])
    if zip_city:
        story.append(Paragraph(zip_city, normal))
    story.append(Spacer(1, 8))

    issue = _pick_issue_date(invoice)
    meta = Table(
        [[Paragraph(f"<b>Rechnung Nummer: {_safe_text(invoice.invoice_number) or '-'}</b>", normal), Paragraph(f"<b>Rechnungsdatum: {_format_date_ddmmyy(issue)}</b>", ParagraphStyle("RightBold", parent=bold, alignment=2))]],
        colWidths=[95 * mm, 80 * mm],
    )
    meta.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0)]))
    story.extend([meta, Spacer(1, 8)])

    site_label, period, bullet, fixed_price = _project_box_rows(invoice)
    project = Table(
        [
            [Paragraph("<b>Baustelle</b>", normal), Paragraph(site_label, normal)],
            [Paragraph("<b>Ausfuehrungszeitraum:</b>", normal), Paragraph(period, normal)],
            [Paragraph(f"{bullet}<br/>{fixed_price}", normal), ""],
        ],
        colWidths=[44 * mm, 131 * mm],
    )
    project.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("SPAN", (0, 2), (1, 2)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([project, Spacer(1, 10)])

    if kind == "pauschal":
        computed = sum(_compute_line_amount(line) for line in invoice.lines or [])
        agreed = float(invoice.pauschal_amount) if invoice.pauschal_amount is not None else None
        amount = agreed if agreed is not None else computed
        pauschal = Table(
            [[Paragraph("<b>Zum vereinbarten Pauschalpreis</b>", normal), Paragraph(f"<b>{_format_money_eur(amount)}</b>", ParagraphStyle("MoneyRight", parent=bold, alignment=2))]],
            colWidths=[120 * mm, 55 * mm],
        )
        pauschal.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0)]))
        story.extend([pauschal, Spacer(1, 8)])
    else:
        rows = [
            [
                Paragraph("<b>Datum</b>", small),
                Paragraph("<b>Beschreibung</b>", small),
                Paragraph("<b>Mitarbeiter</b>", small),
                Paragraph("<b>Auftrag</b>", small),
                Paragraph("<b>Baustelle</b>", small),
                Paragraph("<b>Stunden</b>", ParagraphStyle("R1", parent=small, alignment=2)),
                Paragraph("<b>Satz</b>", ParagraphStyle("R2", parent=small, alignment=2)),
                Paragraph("<b>Betrag</b>", ParagraphStyle("R3", parent=small, alignment=2)),
            ]
        ]
        total_hours = 0.0
        total_amount = 0.0
        for line in _sorted_lines(invoice):
            work_entry = line.work_entry
            hours = float(line.hours_allocated or 0)
            rate = _compute_line_rate(line)
            amount = _compute_line_amount(line)
            total_hours += hours
            total_amount += amount
            rows.append(
                [
                    Paragraph(_format_date_yyyymmdd(line.service_date) if line.service_date else "-", small),
                    Paragraph(_safe_text(line.description) or _safe_text(getattr(work_entry, "description", None)) or "-", small),
                    Paragraph(f"{work_entry.employee.first_name} {work_entry.employee.last_name}" if work_entry and work_entry.employee else "-", small),
                    Paragraph(_safe_text(getattr(getattr(work_entry, "order", None), "title", None)) or "-", small),
                    Paragraph(_safe_text(getattr(getattr(work_entry, "site", None), "site_name", None)) or "-", small),
                    Paragraph(f"{hours:.2f}", ParagraphStyle("Hr", parent=small, alignment=2)),
                    Paragraph(_format_money_ui(rate), ParagraphStyle("Rt", parent=small, alignment=2)),
                    Paragraph(_format_money_ui(amount), ParagraphStyle("Amt", parent=small, alignment=2)),
                ]
            )

        positions = Table(rows, colWidths=[19 * mm, 32 * mm, 26 * mm, 26 * mm, 18 * mm, 15 * mm, 18 * mm, 21 * mm], repeatRows=1)
        positions.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.extend([positions, Spacer(1, 6)])
        totals = Table(
            [
                ["", Paragraph("<b>Summe Stunden:</b>", ParagraphStyle("TR1", parent=normal, alignment=2)), Paragraph(f"{total_hours:.2f}", ParagraphStyle("TR2", parent=normal, alignment=2))],
                ["", Paragraph("<b>Summe Betrag:</b>", ParagraphStyle("TR3", parent=normal, alignment=2)), Paragraph(_format_money_ui(total_amount), ParagraphStyle("TR4", parent=normal, alignment=2))],
            ],
            colWidths=[111 * mm, 32 * mm, 32 * mm],
        )
        totals.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0)]))
        story.extend([totals, Spacer(1, 10)])

    due = issue + timedelta(days=10)
    legal_lines = [
        f"Zahlbar: Rein Nettokasse bis zum {_format_date_ddmmyy(due)} nach Rechnungsstellung.",
        "Gemaess Paragraph 13b UStG geht die Umsatzsteuerschuld auf den Auftraggeber/Rechnungsempfaenger ueber.",
        "Es besteht Steuerschuldnerschaft des Leistungsempfaengers!",
        "Privatpersonen und Unternehmer fuer ihren nichtunternehmerischen Bereich haben eine Rechnungsaufbewahrungspflicht von 2 Jahren.",
        "Wir freuen uns, mit Ihnen gemeinsam arbeiten zu duerfen!",
    ]
    for text in legal_lines:
        story.append(Paragraph(text, normal))
        story.append(Spacer(1, 4))

    doc.build(story, onFirstPage=_add_footer, onLaterPages=_add_footer)
    return buffer.getvalue()


def _set_cell_text(cell, text: str, *, bold: bool = False, align: int | None = None, size: float = 10) -> None:
    from docx.shared import Pt

    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def build_invoice_docx(invoice, kind: str = "detailed") -> bytes:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.75)

    footer = section.footer
    for idx, line in enumerate(FOOTER_LINES):
        p = footer.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.runs:
            p.runs[0].font.size = Pt(8)
        if idx == 0:
            p.paragraph_format.space_before = Pt(6)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.LEFT
    header.autofit = False
    header.columns[0].width = Inches(4.5)
    header.columns[1].width = Inches(3.0)
    _remove_table_borders(header)
    for cell in header.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    left = header.cell(0, 0).paragraphs[0]
    run = left.add_run("KUNDENRECHNUNG")
    run.bold = True
    run.font.size = Pt(16)
    right_p = header.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo = _logo_path()
    if logo:
        right_p.add_run().add_picture(str(logo), width=Inches(1.2))
    for key in ("name", "street", "zipCity", "phone", "email"):
        p = header.cell(0, 1).add_paragraph(SELLER[key])
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(10)

    customer = invoice.customer
    doc.add_paragraph(_safe_text(customer.company_name) or "-")
    if _safe_text(customer.street):
        doc.add_paragraph(customer.street)
    zip_city = " ".join([part for part in [_safe_text(customer.zip_code), _safe_text(customer.city)] if part])
    if zip_city:
        doc.add_paragraph(zip_city)

    meta = doc.add_table(rows=1, cols=2)
    meta.autofit = False
    meta.columns[0].width = Inches(4.5)
    meta.columns[1].width = Inches(3.0)
    _remove_table_borders(meta)
    _set_cell_text(meta.cell(0, 0), f"Rechnung Nummer: {_safe_text(invoice.invoice_number) or '-'}", bold=True)
    _set_cell_text(meta.cell(0, 1), f"Rechnungsdatum: {_format_date_ddmmyy(_pick_issue_date(invoice))}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    site_label, period, bullet, fixed_price = _project_box_rows(invoice)
    project = doc.add_table(rows=3, cols=2)
    project.style = "Table Grid"
    project.autofit = False
    project.columns[0].width = Inches(2.0)
    project.columns[1].width = Inches(5.2)
    _set_cell_text(project.cell(0, 0), "Baustelle", bold=True)
    _set_cell_text(project.cell(0, 1), site_label)
    _set_cell_text(project.cell(1, 0), "Ausfuehrungszeitraum:", bold=True)
    _set_cell_text(project.cell(1, 1), period)
    merged = project.cell(2, 0).merge(project.cell(2, 1))
    _set_cell_text(merged, f"{bullet}\n{fixed_price}")

    if kind == "pauschal":
        computed = sum(_compute_line_amount(line) for line in invoice.lines or [])
        agreed = float(invoice.pauschal_amount) if invoice.pauschal_amount is not None else None
        amount = agreed if agreed is not None else computed
        pauschal = doc.add_table(rows=1, cols=2)
        pauschal.autofit = False
        pauschal.columns[0].width = Inches(5.2)
        pauschal.columns[1].width = Inches(2.0)
        _remove_table_borders(pauschal)
        _set_cell_text(pauschal.cell(0, 0), "Zum vereinbarten Pauschalpreis", bold=True)
        _set_cell_text(pauschal.cell(0, 1), _format_money_eur(amount), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        rows = doc.add_table(rows=1, cols=8)
        rows.style = "Table Grid"
        rows.autofit = False
        widths = [0.75, 1.3, 1.0, 1.0, 0.75, 0.6, 0.8, 0.9]
        titles = ["Datum", "Beschreibung", "Mitarbeiter", "Auftrag", "Baustelle", "Stunden", "Satz", "Betrag"]
        for idx, title in enumerate(titles):
            rows.columns[idx].width = Inches(widths[idx])
            _set_cell_text(rows.cell(0, idx), title, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT if idx >= 5 else None, size=8)
        total_hours = 0.0
        total_amount = 0.0
        for line in _sorted_lines(invoice):
            work_entry = line.work_entry
            hours = float(line.hours_allocated or 0)
            rate = _compute_line_rate(line)
            amount = _compute_line_amount(line)
            total_hours += hours
            total_amount += amount
            cells = rows.add_row().cells
            values = [
                _format_date_yyyymmdd(line.service_date) if line.service_date else "-",
                _safe_text(line.description) or _safe_text(getattr(work_entry, "description", None)) or "-",
                f"{work_entry.employee.first_name} {work_entry.employee.last_name}" if work_entry and work_entry.employee else "-",
                _safe_text(getattr(getattr(work_entry, "order", None), "title", None)) or "-",
                _safe_text(getattr(getattr(work_entry, "site", None), "site_name", None)) or "-",
                f"{hours:.2f}",
                _format_money_ui(rate),
                _format_money_ui(amount),
            ]
            for idx, value in enumerate(values):
                _set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.RIGHT if idx >= 5 else None, size=8)

        totals = doc.add_table(rows=2, cols=3)
        totals.autofit = False
        totals.columns[0].width = Inches(4.8)
        totals.columns[1].width = Inches(1.4)
        totals.columns[2].width = Inches(1.4)
        _remove_table_borders(totals)
        _set_cell_text(totals.cell(0, 1), "Summe Stunden:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(totals.cell(0, 2), f"{total_hours:.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(totals.cell(1, 1), "Summe Betrag:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(totals.cell(1, 2), _format_money_ui(total_amount), align=WD_ALIGN_PARAGRAPH.RIGHT)

    issue = _pick_issue_date(invoice)
    due = issue + timedelta(days=10)
    for text in [
        f"Zahlbar: Rein Nettokasse bis zum {_format_date_ddmmyy(due)} nach Rechnungsstellung.",
        "Gemaess Paragraph 13b UStG geht die Umsatzsteuerschuld auf den Auftraggeber/Rechnungsempfaenger ueber.",
        "Es besteht Steuerschuldnerschaft des Leistungsempfaengers!",
        "Privatpersonen und Unternehmer fuer ihren nichtunternehmerischen Bereich haben eine Rechnungsaufbewahrungspflicht von 2 Jahren.",
        "Wir freuen uns, mit Ihnen gemeinsam arbeiten zu duerfen!",
    ]:
        doc.add_paragraph(text)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

