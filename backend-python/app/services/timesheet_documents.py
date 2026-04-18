from __future__ import annotations

from io import BytesIO
from pathlib import Path




def _logo_path() -> Path | None:
    candidate = Path(__file__).resolve().parents[3] / "backend" / "assets" / "zmd-deco-logo.png"
    return candidate if candidate.exists() else None


def build_timesheet_pdf(data: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=40, rightMargin=40, topMargin=35, bottomMargin=45)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("TimesheetNormal", parent=styles["Normal"], fontName="Helvetica", fontSize=8, leading=10)
    title = ParagraphStyle("TimesheetTitle", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=12)

    story = []
    logo = _logo_path()
    if logo:
        header = Table(
            [
                [
                    "",
                    Image(str(logo), width=24 * mm, height=24 * mm),
                ],
                [
                    "",
                    Paragraph("Z&M Deco Bremen", normal),
                ],
            ],
            colWidths=[125 * mm, 50 * mm],
        )
        header.setStyle(
            TableStyle(
                [
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        story.append(header)

    header_line = Table(
        [
            [
                Paragraph(f"Stunden Zettel: {data['monthName']} {data['year']}", title),
                Paragraph(f"Name: {data['employee']['lastName']}", title),
            ]
        ],
        colWidths=[90 * mm, 85 * mm],
    )
    header_line.setStyle(
        TableStyle(
            [
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([header_line, Spacer(1, 4)])

    table_rows = [
        [
            Paragraph(data["monthName"], normal),
            Paragraph("Arbeitszeit<br/>(Abzueglich Pause)", normal),
            Paragraph("Beginn", normal),
            Paragraph("Ende", normal),
        ]
    ]
    for row in data["rows"]:
        table_rows.append(
            [
                Paragraph(row["dateLabel"], normal),
                Paragraph(row["workLabel"], normal),
                Paragraph(row["begin"], normal),
                Paragraph(row["end"], normal),
            ]
        )

    table = Table(table_rows, colWidths=[46 * mm, 66 * mm, 31 * mm, 31 * mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([table, Spacer(1, 8)])
    story.append(Paragraph(f"<b>Gesamtstunden: {data['totalHoursLabel']} Std</b>", normal))
    signs = Table([["Arbeitsnehmer:", "Arbeitsgeber:"]], colWidths=[87 * mm, 87 * mm])
    signs.setStyle(
        TableStyle(
            [
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    story.append(signs)
    doc.build(story)
    return buffer.getvalue()


def build_timesheet_docx(data: dict) -> bytes:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.LEFT
    top.autofit = False
    top.columns[0].width = Inches(4.9)
    top.columns[1].width = Inches(2.8)
    for cell in top.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    logo = _logo_path()
    if logo:
        run = top.cell(0, 1).paragraphs[0].add_run()
        run.add_picture(str(logo), width=Inches(1.0))
    top.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = top.cell(0, 1).add_paragraph("Z&M Deco Bremen")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(10)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Inches(4.9)
    header.columns[1].width = Inches(2.8)
    left = header.cell(0, 0).paragraphs[0]
    left.add_run(f"Stunden Zettel: {data['monthName']} {data['year']}").bold = True
    right = header.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"Name: {data['employee']['lastName']}").bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.8), Inches(2.6), Inches(1.2), Inches(1.2)]
    titles = [data["monthName"], "Arbeitszeit\n(Abzueglich Pause)", "Beginn", "Ende"]
    for idx, title_text in enumerate(titles):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        para = cell.paragraphs[0]
        para.add_run(title_text)
        para.runs[0].font.size = Pt(9)
    for row in data["rows"]:
        cells = table.add_row().cells
        values = [row["dateLabel"], row["workLabel"], row["begin"], row["end"]]
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            run = cells[idx].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)

    total = doc.add_paragraph()
    total.add_run(f"Gesamtstunden: {data['totalHoursLabel']} Std").bold = True
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.autofit = False
    sign_table.columns[0].width = Inches(3.7)
    sign_table.columns[1].width = Inches(3.7)
    sign_table.cell(0, 0).paragraphs[0].add_run("Arbeitsnehmer:")
    p2 = sign_table.cell(0, 1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run("Arbeitsgeber:")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

